
from copy import deepcopy
from typing import List, Tuple

from docx import Document
from docx.shared import Pt


def fill_full_page(
    template_doc: Document,
    id_pairs: List[Tuple[str, str, str]],
    *,
    sets_per_page: int = 8,
) -> Document:
    """
    Fill a DOCX template with up to `sets_per_page` sets of IDs (RID, MID, Visit).

    Placeholders are expected to be in the form:
    - {RID 1}, {MID 1}, {V 1}
    - ...
    - {RID N}, {MID N}, {V N}
    """
    page = deepcopy(template_doc)

    def replace_placeholders(paragraphs, rid: str, mid: str, visit: str, idx: int) -> None:
        target_rid = f"{{RID {idx}}}"
        target_mid = f"{{MID {idx}}}"
        target_v = f"{{V {idx}}}"

        for para in paragraphs:
            full_text = "".join(run.text for run in para.runs)
            if target_rid in full_text or target_mid in full_text or target_v in full_text:
                full_text = (
                    full_text.replace(target_rid, str(rid))
                    .replace(target_mid, str(mid))
                    .replace(target_v, str(visit))
                )

                for run in para.runs:
                    run.text = ""

                if para.runs:
                    para.runs[0].text = full_text
                    font = para.runs[0].font
                    font.name = "Calibri"
                    font.size = Pt(11)

    for idx in range(1, sets_per_page + 1):
        rid, mid, visit = ("", "", "")
        if idx <= len(id_pairs):
            rid, mid, visit = id_pairs[idx - 1]

        replace_placeholders(page.paragraphs, rid, mid, visit, idx)

        for table in page.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholders(cell.paragraphs, rid, mid, visit, idx)

    return page

